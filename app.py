"""
Video Lecture Summarizer - Complete Application
===============================================
Extract summaries from video lectures using AI

Features:
- Video URL input (YouTube, Vimeo, etc.)
- Speech-to-text transcription
- Dual models: TF-IDF (Extractive) & BART (Abstractive)
- Beautiful UI with statistics
- Download summaries as TXT/DOCX

Author: Your Name
"""

import streamlit as st
import pickle
import numpy as np
import pandas as pd
from pathlib import Path
import io
from datetime import datetime
import tempfile
import os
import time

# Page configuration
st.set_page_config(
    page_title="AI Lecture Summarizer",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for beautiful UI
st.markdown("""
<style>
    /* Main container */
    .main {
        padding: 2rem;
    }
    
    /* Header styling */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .main-header h1 {
        color: white;
        font-size: 2.5rem;
        font-weight: bold;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
    }
    
    .main-header p {
        color: rgba(255,255,255,0.9);
        font-size: 1.2rem;
        margin-top: 0.5rem;
    }
    
    /* Pipeline steps */
    .pipeline-step {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
        text-align: center;
        font-weight: bold;
        font-size: 1.1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    /* Cards */
    .stat-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    
    .model-card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border: 2px solid #667eea;
    }
    
    /* Buttons */
    .stButton>button {
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        font-weight: bold;
        padding: 0.75rem 2rem;
        border-radius: 8px;
        border: none;
        font-size: 1.1rem;
        transition: transform 0.2s;
    }
    
    .stButton>button:hover {
        transform: scale(1.05);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }
    
    /* Success/Info boxes */
    .success-box {
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    
    .info-box {
        background-color: #d1ecf1;
        border-left: 5px solid #17a2b8;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    
    .warning-box {
        background-color: #fff3cd;
        border-left: 5px solid #ffc107;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    
    /* Progress indicators */
    .stProgress > div > div {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Metric styling */
    [data-testid="stMetricValue"] {
        font-size: 2rem;
        color: #667eea;
        font-weight: bold;
    }
    
    /* Summary box */
    .summary-box {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid #667eea;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# STEP 1: VIDEO → AUDIO
# ==============================================================================

@st.cache_data(show_spinner=False)
def extract_audio_from_url(video_url):
    """Extract audio from video URL - Streamlit Cloud compatible"""
    try:
        import yt_dlp
        import glob
        
        temp_dir = tempfile.mkdtemp()
        output_template = os.path.join(temp_dir, 'audio')
        
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': output_template,
            'quiet': False,  # Show errors
            'no_warnings': False,
            'extract_audio': True,
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '128',
            }],
            # Streamlit Cloud compatibility
            'socket_timeout': 30,
            'retries': 3,
            # Bypass YouTube bot detection
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-us,en;q=0.5',
                'Sec-Fetch-Mode': 'navigate',
            },
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(video_url, download=True)
            title = info.get('title', 'Unknown')
            duration = info.get('duration', 0)
        
        # Find the downloaded file
        audio_files = glob.glob(os.path.join(temp_dir, 'audio*'))
        if not audio_files:
            st.error("Audio file not created after download")
            return None, None, None
        
        return audio_files[0], title, duration
    
    except Exception as e:
        st.error(f"Audio extraction error: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None, None, None

# ==============================================================================
# STEP 2: AUDIO → TRANSCRIPT
# ==============================================================================

@st.cache_data(show_spinner=False)
def transcribe_audio_whisper(audio_path):
    """Transcribe audio using Whisper"""
    try:
        import whisper
        
        model = whisper.load_model("base")
        result = model.transcribe(audio_path)
        
        return result['text'], result.get('segments', [])
    
    except Exception as e:
        return None, None

# ==============================================================================
# STEP 3: LOAD MODELS
# ==============================================================================

@st.cache_resource
def load_tfidf_model(model_path='tfidf_lecture_model.pkl'):
    """Load TF-IDF model"""
    try:
        with open(model_path, 'rb') as f:
            saved = pickle.load(f)
        return saved['model'], saved['vectorizer']
    except:
        return None, None

@st.cache_resource
def load_bart_model(model_name='./bart_lecture_model'):
    """Load BART model from local folder"""
    try:
        from transformers import BartForConditionalGeneration, BartTokenizer
        import torch
        
        model = BartForConditionalGeneration.from_pretrained(model_name)
        tokenizer = BartTokenizer.from_pretrained(model_name)
        
        model.eval()
        
        return model, tokenizer
    
    except Exception as e:
        st.error(f"Error loading BART: {str(e)}")
        return None, None

# ==============================================================================
# STEP 4: SUMMARIZATION
# ==============================================================================

def split_into_sentences(text):
    """Simple sentence splitter"""
    sentences = []
    for sent in text.split('.'):
        sent = sent.strip()
        if sent and len(sent) > 5:
            sentences.append(sent + '.')
    return sentences

def chunk_text(text, max_tokens=900):
    """
    Split long text into chunks that fit within token limit.
    Uses sentence boundaries to avoid cutting mid-sentence.
    """
    sentences = split_into_sentences(text)
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        # Rough estimate: 1 token ≈ 4 characters
        sentence_tokens = len(sentence) // 4
        
        if current_length + sentence_tokens > max_tokens:
            # Save current chunk and start new one
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_length = sentence_tokens
        else:
            current_chunk.append(sentence)
            current_length += sentence_tokens
    
    # Add last chunk
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def summarize_with_tfidf(text, model, vectorizer, summary_ratio=0.2):
    """Generate summary using TF-IDF - handles any length!"""
    sentences = split_into_sentences(text)
    
    if len(sentences) == 0:
        return "No sentences found.", pd.DataFrame(), {}
    
    # Show info for very long documents
    if len(sentences) > 500:
        st.info(f"📄 Long transcript detected! Analyzing {len(sentences)} sentences...")
    
    # Vectorize and predict (TF-IDF can handle any length!)
    X = vectorizer.transform(sentences)
    predictions = model.predict(X)
    probabilities = model.predict_proba(X)[:, 1]
    
    # Create results
    results_df = pd.DataFrame({
        'sentence': sentences,
        'prediction': predictions,
        'probability': probabilities,
        'index': range(len(sentences))
    })
    
    # Get top sentences
    target_sentences = max(1, int(len(sentences) * summary_ratio))
    important = results_df.nlargest(target_sentences, 'probability')
    important = important.sort_values('index')
    
    # Generate summary
    summary = ' '.join(important['sentence'].tolist())
    
    # Statistics
    stats = {
        'original_sentences': len(sentences),
        'summary_sentences': len(important),
        'compression_ratio': (len(important) / len(sentences)) * 100,
        'avg_confidence': important['probability'].mean(),
        'model_type': 'Extractive (TF-IDF)'
    }
    
    return summary, results_df, stats

def summarize_with_bart(text, model, tokenizer, summary_ratio=0.2):
    """
    Generate summary using BART with automatic chunking for long documents.
    Handles transcripts of any length!
    """
    import torch
    
    # Split long text into manageable chunks
    chunks = chunk_text(text, max_tokens=900)
    
    total_sentences = len(split_into_sentences(text))
    chunk_summaries = []
    
    # Show progress for long documents
    if len(chunks) > 1:
        st.info(f"📄 Long transcript detected! Processing in {len(chunks)} chunks...")
    
    # Summarize each chunk
    progress_bar = st.progress(0) if len(chunks) > 3 else None
    
    for i, chunk in enumerate(chunks):
        # Calculate target length for this chunk
        chunk_words = len(chunk.split())
        target_length = max(30, int(chunk_words * summary_ratio * 1.5))  # Slightly longer per chunk
        
        # Tokenize
        inputs = tokenizer(
            chunk,
            max_length=1024,
            truncation=True,
            return_tensors="pt"
        )
        
        # Generate summary for this chunk
        with torch.no_grad():
            summary_ids = model.generate(
                inputs["input_ids"],
                max_length=target_length,
                min_length=max(20, target_length // 3),
                length_penalty=2.0,
                num_beams=4,
                early_stopping=True,
                no_repeat_ngram_size=3
            )
        
        # Decode
        chunk_summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        chunk_summaries.append(chunk_summary)
        
        # Update progress
        if progress_bar:
            progress_bar.progress((i + 1) / len(chunks))
    
    if progress_bar:
        progress_bar.empty()
    
    # Combine chunk summaries
    if len(chunks) == 1:
        # Single chunk - use as is
        final_summary = chunk_summaries[0]
    else:
        # Multiple chunks - combine and optionally re-summarize
        combined = ' '.join(chunk_summaries)
        
        # If combined summary is still very long, do final pass
        if len(combined.split()) > 500:
            st.info("🔄 Performing final consolidation pass...")
            
            # Final summarization pass
            final_inputs = tokenizer(
                combined,
                max_length=1024,
                truncation=True,
                return_tensors="pt"
            )
            
            target_final = max(100, int(len(combined.split()) * 0.6))
            
            with torch.no_grad():
                final_ids = model.generate(
                    final_inputs["input_ids"],
                    max_length=target_final,
                    min_length=max(50, target_final // 2),
                    length_penalty=2.0,
                    num_beams=4,
                    early_stopping=True,
                    no_repeat_ngram_size=3
                )
            
            final_summary = tokenizer.decode(final_ids[0], skip_special_tokens=True)
        else:
            final_summary = combined
    
    # Statistics
    summary_sentences = len(split_into_sentences(final_summary))
    
    stats = {
        'original_sentences': total_sentences,
        'summary_sentences': summary_sentences,
        'compression_ratio': (len(final_summary) / len(text)) * 100,
        'chunks_processed': len(chunks),
        'model_type': 'Abstractive (BART with Chunking)'
    }
    
    if len(chunks) > 1:
        st.success(f"✅ Successfully processed {len(chunks)} chunks into final summary!")
    
    return final_summary, pd.DataFrame(), stats

# ==============================================================================
# STEP 5: DOWNLOAD FUNCTIONS
# ==============================================================================

def create_txt_download(summary, transcript, stats, video_title=""):
    """Create downloadable TXT file"""
    content = f"""VIDEO LECTURE SUMMARY
{'='*80}
Video: {video_title}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Model: {stats.get('model_type', 'N/A')}

STATISTICS:
- Original sentences: {stats.get('original_sentences', 'N/A')}
- Summary sentences: {stats.get('summary_sentences', 'N/A')}
- Compression ratio: {stats.get('compression_ratio', 0):.1f}%

SUMMARY:
{'='*80}
{summary}

{'='*80}
FULL TRANSCRIPT:
{'='*80}
{transcript}
"""
    return content

def create_docx_download(summary, transcript, stats, video_title=""):
    """Create downloadable DOCX file"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Video Lecture Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Video info
        if video_title:
            video_para = doc.add_paragraph(f"Video: {video_title}")
            video_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        model_para = doc.add_paragraph(f"Model: {stats.get('model_type', 'N/A')}")
        model_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Statistics
        doc.add_heading('Statistics', 1)
        doc.add_paragraph(f"Original sentences: {stats.get('original_sentences', 'N/A')}")
        doc.add_paragraph(f"Summary sentences: {stats.get('summary_sentences', 'N/A')}")
        doc.add_paragraph(f"Compression ratio: {stats.get('compression_ratio', 0):.1f}%")
        
        doc.add_paragraph()
        
        # Summary
        doc.add_heading('Summary', 1)
        doc.add_paragraph(summary)
        
        doc.add_page_break()
        
        # Full transcript
        doc.add_heading('Full Transcript', 1)
        doc.add_paragraph(transcript)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    except ImportError:
        return None

# ==============================================================================
# MAIN APP
# ==============================================================================

def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🎓 AI Lecture Summarizer</h1>
        <p>Transform video lectures into concise summaries using AI</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar - Settings
    st.sidebar.title("⚙️ Configuration")
    
    # Model selection
    st.sidebar.markdown("### 🤖 Choose AI Model")
    model_choice = st.sidebar.radio(
        "Select summarization model:",
        ["TF-IDF + Random Forest (Extractive)", "BART Fine-tuned (Abstractive)"],
        help="Extractive: Selects important sentences\nAbstractive: Generates new paraphrased text"
    )
    
    # Summary length
    st.sidebar.markdown("### 📏 Summary Length")
    summary_ratio = st.sidebar.slider(
        "Target summary size:",
        min_value=10,
        max_value=50,
        value=20,
        step=5,
        help="Percentage of original content to keep",
        format="%d%%"
    ) / 100
    
    st.sidebar.markdown("---")
    
    # Instructions
    with st.sidebar.expander("📖 How to Use"):
        st.markdown("""
        1. **Enter video URL** (YouTube, Vimeo, etc.)
        2. **Click Process** to start
        3. **Choose your model** (TF-IDF or BART)
        4. **Adjust summary length** if needed
        5. **Download** your summary
        
        **Models:**
        - **TF-IDF:** Fast, selects key sentences
        - **BART:** Smart, creates new summaries
        """)
    
    with st.sidebar.expander("🎯 Supported Platforms"):
        st.markdown("""
        **Best for Cloud Deployment:**
        ✅ Archive.org (Recommended)
        ✅ Vimeo  
        ✅ Dailymotion  
        
        **May have restrictions:**
        ⚠️ YouTube (may be blocked on cloud)
        
        **Tip:** Use Manual Transcript for guaranteed results!
        """)
    
    with st.sidebar.expander("💡 Sample Videos"):
        st.markdown("""
        **Try these (Archive.org):**
        
        MIT Lecture:
        `https://archive.org/details/mit-6.034-fall-2010-lecture-01`
        
        Stanford ML:
        `https://archive.org/details/stanford-machine-learning-cs229`
        """)
    
    # Main content
    st.markdown("---")
    
    # Video URL input
    st.markdown("### 🔗 Enter Video URL")
    
    col1, col2 = st.columns([4, 1])
    
    with col1:
        video_url = st.text_input(
            "Video URL",
            placeholder="https://www.youtube.com/watch?v=...",
            label_visibility="collapsed"
        )
    
    with col2:
        process_button = st.button("🚀 Process Video", use_container_width=True, type="primary")
    
    # Alternative: Manual transcript
    with st.expander("💡 Or paste transcript directly"):
        manual_transcript = st.text_area(
            "Paste your lecture transcript here:",
            height=150,
            placeholder="Enter lecture transcript text..."
        )
        use_manual = st.button("📝 Use Manual Transcript", use_container_width=True)
    
    st.markdown("---")
    
    # Processing pipeline
    if (process_button and video_url) or (use_manual and manual_transcript):
        
        transcript = None
        video_title = "Manual Input"
        duration = 0
        
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        if not use_manual:
            # Step 1: Extract audio
            st.markdown('<div class="pipeline-step">STEP 1: Extracting Audio 🎵</div>', unsafe_allow_html=True)
            status_text.text("⏳ Downloading audio from video...")
            progress_bar.progress(10)
            
            audio_path, video_title, duration = extract_audio_from_url(video_url)
            
            if audio_path:
                status_text.text(f"✅ Audio extracted: {video_title}")
                progress_bar.progress(30)
                time.sleep(0.5)
                
                # Step 2: Transcribe
                st.markdown('<div class="pipeline-step">STEP 2: Transcribing Audio 🎤→📝</div>', unsafe_allow_html=True)
                status_text.text("⏳ Converting speech to text using Whisper AI...")
                progress_bar.progress(40)
                
                transcript, segments = transcribe_audio_whisper(audio_path)
                
                # Cleanup
                if os.path.exists(audio_path):
                    os.remove(audio_path)
                
                if transcript:
                    status_text.text("✅ Transcription complete!")
                    progress_bar.progress(60)
                else:
                    st.error("❌ Transcription failed. Please try manual input.")
                    return
            else:
                st.error("❌ Failed to extract audio. Please check the URL or try manual input.")
                return
        else:
            transcript = manual_transcript
            status_text.text("✅ Using manual transcript")
            progress_bar.progress(60)
        
        if transcript:
            # Show transcript preview
            with st.expander("📄 View Full Transcript", expanded=False):
                st.text_area("Transcript", transcript, height=200, disabled=True, label_visibility="collapsed")
            
            # Step 3: Summarize
            st.markdown('<div class="pipeline-step">STEP 3: Generating Summary 🤖✨</div>', unsafe_allow_html=True)
            status_text.text("⏳ AI is analyzing and summarizing...")
            progress_bar.progress(70)
            
            # Load appropriate model
            if "TF-IDF" in model_choice:
                tfidf_model, tfidf_vectorizer = load_tfidf_model()
                if tfidf_model and tfidf_vectorizer:
                    summary, details, stats = summarize_with_tfidf(
                        transcript, tfidf_model, tfidf_vectorizer, summary_ratio
                    )
                else:
                    st.error("❌ TF-IDF model not found. Please ensure tfidf_lecture_model.pkl is in the same folder.")
                    return
            else:
                bart_model, bart_tokenizer = load_bart_model('./bart_lecture_model')
                if bart_model and bart_tokenizer:
                    summary, details, stats = summarize_with_bart(
                        transcript, bart_model, bart_tokenizer, summary_ratio
                    )
                else:
                    st.error("❌ BART model not found. Check Hugging Face model name.")
                    return
            
            status_text.text("✅ Summary generated!")
            progress_bar.progress(100)
            time.sleep(0.5)
            
            # Clear progress indicators
            progress_bar.empty()
            status_text.empty()
            
            # Display results
            st.markdown("---")
            st.markdown("## 📊 Results")
            
            # Statistics cards
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric(
                    "Original Sentences",
                    stats['original_sentences'],
                    help="Total sentences in transcript"
                )
            
            with col2:
                st.metric(
                    "Summary Sentences",
                    stats['summary_sentences'],
                    help="Sentences in generated summary"
                )
            
            with col3:
                st.metric(
                    "Compression",
                    f"{stats['compression_ratio']:.1f}%",
                    help="Summary size vs original"
                )
            
            with col4:
                if 'avg_confidence' in stats:
                    st.metric(
                        "Avg Confidence",
                        f"{stats.get('avg_confidence', 0):.0%}",
                        help="Model's average confidence"
                    )
                else:
                    st.metric(
                        "Model Type",
                        stats['model_type'].split()[0],
                        help="Summarization method used"
                    )
            
            st.markdown("---")
            
            # Summary display
            st.markdown("### ✨ Generated Summary")
            st.markdown(f"""
            <div class="summary-box">
                <p style="font-size: 1.1rem; line-height: 1.8; color: #2c3e50;">
                    {summary}
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # Model info
            st.info(f"**Model Used:** {stats['model_type']}")
            
            # Download section
            st.markdown("---")
            st.markdown("### 💾 Download Summary")
            
            col1, col2 = st.columns(2)
            
            with col1:
                txt_content = create_txt_download(summary, transcript, stats, video_title)
                st.download_button(
                    label="📄 Download as TXT",
                    data=txt_content,
                    file_name=f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            with col2:
                docx_content = create_docx_download(summary, transcript, stats, video_title)
                if docx_content:
                    st.download_button(
                        label="📝 Download as DOCX",
                        data=docx_content,
                        file_name=f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            # Detailed analysis (for TF-IDF)
            if "TF-IDF" in model_choice and len(details) > 0:
                with st.expander("📈 Detailed Sentence Analysis"):
                    top_sentences = details.nlargest(10, 'probability')[['sentence', 'probability', 'prediction']]
                    top_sentences['prediction'] = top_sentences['prediction'].map({0: '❌ Not Important', 1: '✅ Important'})
                    top_sentences['probability'] = top_sentences['probability'].apply(lambda x: f"{x:.2%}")
                    top_sentences.columns = ['Sentence', 'Confidence', 'Classification']
                    
                    st.dataframe(top_sentences, use_container_width=True, hide_index=True)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #7f8c8d; padding: 2rem;'>
        <p><strong>🎓 AI Lecture Summarizer</strong></p>
        <p>Powered by TF-IDF, BART, and Whisper AI</p>
        <p style='font-size: 0.9rem;'>Built with Streamlit • Machine Learning • NLP</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
